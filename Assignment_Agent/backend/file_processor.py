"""
File processing service to read content from various file types
"""
import os
import json
import csv
from pathlib import Path
from typing import Dict, List
import os
import io
import base64
import requests

# Optional imports for different file types
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional alternate DOCX extraction libraries
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except Exception:
    DOCX2TXT_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except Exception:
    MAMMOTH_AVAILABLE = False

# Optional Windows COM automation for legacy .doc
try:
    import win32com.client  # type: ignore
    WIN32COM_AVAILABLE = True
except Exception:
    WIN32COM_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

import zipfile
import io as _io

# Optional OCR for scanned PDFs
try:
    from pdf2image import convert_from_path
    from PIL import Image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False


class FileProcessor:
    """Process and extract text from various file types"""
    
    # Supported text file extensions
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h',
        '.html', '.css', '.scss', '.json', '.xml', '.yaml', '.yml', '.sql', '.sh', '.bat',
        '.ps1', '.rb', '.php', '.go', '.rs', '.swift', '.kt', '.dart', '.r', '.m', '.pl',
        '.lua', '.scala', '.clj', '.hs', '.elm', '.ex', '.exs', '.erl', '.ml', '.fs',
        '.cs', '.vb', '.asm', '.s', '.asmx', '.vue', '.svelte', '.tsx', '.jsx'
    }

    @staticmethod
    def _call_nvidia_ocr(image_bytes: bytes) -> str:
        """Send image bytes to configured NVIDIA OCR endpoint and return extracted text if available."""
        nv_key = os.getenv('NVIDIA_OCR_API_KEY')
        nv_url = os.getenv('NVIDIA_OCR_URL')
        if not nv_key or not nv_url:
            return None
        try:
            img_b64 = base64.b64encode(image_bytes).decode('utf-8')
            payload = {"image": img_b64}
            headers = {
                "Authorization": f"Bearer {nv_key}",
                "Content-Type": "application/json",
            }
            resp = requests.post(nv_url, json=payload, headers=headers, timeout=60)
            if resp.status_code == 200:
                data = resp.json()
                text = data.get('text') or data.get('result') or ''
                if text and text.strip():
                    return text
        except Exception:
            pass
        return None

    @staticmethod
    def force_ocr(file_path: str) -> str:
        """Force OCR extraction for a file using NVIDIA OCR (if available) or local pytesseract as fallback.

        Returns extracted text or None if nothing found.
        """
        file_path_obj = Path(file_path)
        ext = file_path_obj.suffix.lower()

        if ext == '.pdf':
            # Use pdf2image to convert pages to images and then call NVIDIA OCR or tesseract
            if not PDF2IMAGE_AVAILABLE:
                return None
            try:
                poppler_path = os.getenv('POPPLER_PATH') or None
                images = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)
                parts = []
                for img in images:
                    # Try NVIDIA OCR first
                    try:
                        buf = io.BytesIO()
                        img.save(buf, format='PNG')
                        bytes_data = buf.getvalue()
                        text = FileProcessor._call_nvidia_ocr(bytes_data)
                        if text:
                            parts.append(text)
                            continue
                    except Exception:
                        pass

                    # Fallback to local pytesseract
                    if PYTESSERACT_AVAILABLE:
                        try:
                            parts.append(pytesseract.image_to_string(img))
                        except Exception:
                            pass

                if parts:
                    joined = '\n\n'.join(p for p in parts if p and p.strip())
                    return joined if joined.strip() else None
            except Exception:
                return None
            return None

        if ext == '.docx':
            return FileProcessor._ocr_docx_images(file_path)

        if ext == '.doc':
            # Try to convert .doc to images and OCR them
            return FileProcessor._ocr_doc_images(file_path)

        # For other file types, no OCR performed
        return None
    
    @staticmethod
    def read_file(file_path: str) -> Dict[str, any]:
        """
        Read file content based on file extension
        Returns dict with filename, content, and file_type
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        filename = file_path_obj.name
        extension = file_path_obj.suffix.lower()
        
        try:
            # Text files
            if extension in FileProcessor.TEXT_EXTENSIONS:
                content = FileProcessor._read_text_file(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'text',
                    'extension': extension
                }
            
            # PDF files
            elif extension == '.pdf':
                content = FileProcessor._read_pdf(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'pdf',
                    'extension': extension
                }
            
            # Word .docx
            elif extension == '.docx':
                content = FileProcessor._read_docx(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'docx',
                    'extension': extension
                }
            # Legacy Word .doc (Windows COM fallback)
            elif extension == '.doc':
                content = FileProcessor._read_doc(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'doc',
                    'extension': extension
                }
            
            # Excel files
            elif extension in ['.xlsx', '.xls']:
                content = FileProcessor._read_excel(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'excel',
                    'extension': extension
                }
            
            # CSV files
            elif extension == '.csv':
                content = FileProcessor._read_csv(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'csv',
                    'extension': extension
                }
            
            # JSON files
            elif extension == '.json':
                content = FileProcessor._read_json(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'json',
                    'extension': extension
                }
            
            # PowerPoint files
            elif extension in ['.ppt', '.pptx', '.pptm']:
                content = FileProcessor._read_ppt(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'file_type': 'ppt',
                    'extension': extension
                }
            
            # Default: try to read as text
            else:
                try:
                    content = FileProcessor._read_text_file(file_path)
                    return {
                        'filename': filename,
                        'content': content,
                        'file_type': 'text',
                        'extension': extension
                    }
                except:
                    return {
                        'filename': filename,
                        'content': f"[Binary file - {extension} - Cannot read as text]",
                        'file_type': 'binary',
                        'extension': extension
                    }
        
        except Exception as e:
            return {
                'filename': filename,
                'content': f"[Error reading file: {str(e)}]",
                'file_type': 'error',
                'extension': extension
            }
    
    @staticmethod
    def _read_text_file(file_path: str) -> str:
        """Read text file with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors='replace'
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')
    
    @staticmethod
    def _read_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        if PDFPLUMBER_AVAILABLE:
            try:
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                extracted = '\n\n'.join(text_content)
                if extracted.strip():
                    return extracted
            except Exception as e:
                pass
        
        if PDF_AVAILABLE:
            try:
                text_content = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                extracted = '\n\n'.join(text_content)
                if extracted.strip():
                    return extracted
            except Exception as e:
                return f"[Error reading PDF: {str(e)}]"
        
        # If we reach here, try OCR fallback for scanned PDFs
        if PDF2IMAGE_AVAILABLE:
            try:
                poppler_path = os.getenv('POPPLER_PATH') or None
                images = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)
                ocr_text_parts = []
                for img in images:
                    text = None
                    # Prefer NVIDIA OCR if configured
                    try:
                        nv_key = os.getenv('NVIDIA_OCR_API_KEY')
                        nv_url = os.getenv('NVIDIA_OCR_URL')
                        if nv_key and nv_url:
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            img_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
                            payload = {"image": img_b64}
                            headers = {
                                "Authorization": f"Bearer {nv_key}",
                                "Content-Type": "application/json",
                            }
                            resp = requests.post(nv_url, json=payload, headers=headers, timeout=60)
                            if resp.status_code == 200:
                                data = resp.json()
                                text = data.get('text') or data.get('result') or ''
                    except Exception:
                        text = None
                    # Fallback to local Tesseract if NVIDIA not configured or failed
                    if (not text or not str(text).strip()) and PYTESSERACT_AVAILABLE:
                        try:
                            text = pytesseract.image_to_string(img)
                        except Exception:
                            text = None
                    if text:
                        ocr_text_parts.append(str(text))
                ocr_text = '\n\n'.join(ocr_text_parts)
                if ocr_text.strip():
                    return ocr_text
            except Exception:
                pass

        # Libraries not available or OCR failed
        if not (PDFPLUMBER_AVAILABLE or PDF_AVAILABLE):
            return "[PDF library not available. Install PyPDF2 or pdfplumber]"
        if not PDF2IMAGE_AVAILABLE:
            return "[No extractable text found. Enable OCR by installing pdf2image & pillow. For local OCR also install pytesseract and set POPPLER_PATH on Windows]"
        return "[No extractable text found in PDF]"
    
    @staticmethod
    def _ocr_docx_images(file_path: str) -> str:
        """Extract images from .docx and run OCR on them using NVIDIA OCR (if configured) or pytesseract."""
        try:
            texts = []
            with zipfile.ZipFile(file_path, 'r') as z:
                for name in z.namelist():
                    if name.startswith('word/media/'):
                        try:
                            data = z.read(name)
                            # Try NVIDIA OCR first
                            try:
                                nvinf = FileProcessor._call_nvidia_ocr(data)
                                if nvinf and nvinf.strip():
                                    texts.append(nvinf.strip())
                                    continue
                            except Exception:
                                pass

                            # Fallback to local tesseract if available
                            if PYTESSERACT_AVAILABLE:
                                try:
                                    from PIL import Image
                                    img = Image.open(_io.BytesIO(data))
                                    text = pytesseract.image_to_string(img)
                                    if text and text.strip():
                                        texts.append(text.strip())
                                except Exception:
                                    continue
                        except Exception:
                            continue
            if texts:
                return '\n\n'.join(texts)
        except Exception:
            return None
        return None

    @staticmethod
    def _ocr_doc_images(file_path: str) -> str:
        """Convert .doc to PDF using win32com, then convert PDF to images and run OCR."""
        # Check prerequisites
        if not WIN32COM_AVAILABLE:
            return None
        
        # If PDF2IMAGE not available, we can still try to extract text via COM and return it
        # This provides a fallback even without OCR capabilities
        if not PDF2IMAGE_AVAILABLE:
            # Try one more time with COM to extract text
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(Path(file_path).absolute()))
                try:
                    # Try different methods to extract text
                    text = doc.Content.Text
                    if isinstance(text, str) and text.strip():
                        return text
                    # Try paragraphs
                    paragraphs = []
                    for para in doc.Paragraphs:
                        para_text = para.Range.Text
                        if para_text and para_text.strip():
                            paragraphs.append(para_text.strip())
                    if paragraphs:
                        return '\n\n'.join(paragraphs)
                finally:
                    doc.Close(False)
                    word.Quit()
            except Exception:
                pass
            return None
        
        try:
            import tempfile
            import shutil
            
            # Create temporary directory for PDF
            temp_dir = tempfile.mkdtemp()
            pdf_path = os.path.join(temp_dir, "temp_doc.pdf")
            
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(Path(file_path).absolute()))
                
                try:
                    # Export document as PDF
                    doc.ExportAsFixedFormat(
                        OutputFileName=pdf_path,
                        ExportFormat=17,  # wdExportFormatPDF
                        OpenAfterExport=False,
                        OptimizeFor=0,  # wdExportOptimizeForPrint
                        BitmapMissingFonts=True,
                        DocStructureTags=True,
                        CreateBookmarks=0,
                        UseISO19005_1=False
                    )
                    
                    # Verify PDF was created
                    if not os.path.exists(pdf_path):
                        return None
                    
                    # Convert PDF pages to images and OCR
                    poppler_path = os.getenv('POPPLER_PATH') or None
                    images = convert_from_path(pdf_path, dpi=200, poppler_path=poppler_path)
                    texts = []
                    
                    for img in images:
                        # Try NVIDIA OCR first
                        try:
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            bytes_data = buf.getvalue()
                            text = FileProcessor._call_nvidia_ocr(bytes_data)
                            if text and text.strip():
                                texts.append(text)
                                continue
                        except Exception:
                            pass
                        
                        # Fallback to local pytesseract
                        if PYTESSERACT_AVAILABLE:
                            try:
                                text = pytesseract.image_to_string(img)
                                if text and text.strip():
                                    texts.append(text)
                            except Exception:
                                pass
                    
                    if texts:
                        return '\n\n'.join(texts)
                        
                finally:
                    doc.Close(False)
                    word.Quit()
            finally:
                # Clean up temp directory
                try:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                except Exception:
                    pass
        except Exception as e:
            # Log error for debugging but return None silently
            import logging
            logger = logging.getLogger(__name__)
            logger.debug(f"OCR for DOC file failed: {e}")
            return None
        return None

    def _read_docx(file_path: str) -> str:
        """Extract text from Word document, including paragraphs, tables, headers and footers.

        This function tries multiple strategies in order:
        1. python-docx structured extraction (paragraphs, tables, headers/footers)
        2. docx2txt extraction (if installed)
        3. mammoth extraction (if installed)
        4. OCR images embedded in the .docx (if pytesseract is installed)
        """
        if not DOCX_AVAILABLE:
            # Fall back to docx2txt or mammoth if python-docx not installed
            if DOCX2TXT_AVAILABLE:
                try:
                    text = docx2txt.process(file_path)
                    if text and text.strip():
                        return text
                except Exception:
                    pass
            if MAMMOTH_AVAILABLE:
                try:
                    with open(file_path, 'rb') as f:
                        result = mammoth.extract_raw_text(f)
                        text = (result.value or '').strip() if result else ''
                        if text:
                            return text
                except Exception:
                    pass
            return "[python-docx library not available]"

        try:
            doc = Document(file_path)
            parts = []

            # Paragraphs
            for para in doc.paragraphs:
                text = (para.text or "").strip()
                if text:
                    parts.append(text)

            # Tables (preserve cell text in readable rows)
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [ (cell.text or "").strip() for cell in row.cells ]
                    table_rows.append('\t'.join(cells))
                if table_rows:
                    parts.append("Table:\n" + "\n".join(table_rows))

            # Headers and Footers
            try:
                for section in doc.sections:
                    header = section.header
                    for para in header.paragraphs:
                        h = (para.text or "").strip()
                        if h:
                            parts.append("Header: " + h)
                    footer = section.footer
                    for para in footer.paragraphs:
                        f = (para.text or "").strip()
                        if f:
                            parts.append("Footer: " + f)
            except Exception:
                # Some docs may not expose section header/footer reliably; ignore failures
                pass

            extracted = "\n\n".join(parts).strip()

            # If nothing found, try alternate extractors
            if not extracted:
                # Try docx2txt
                if DOCX2TXT_AVAILABLE:
                    try:
                        t = docx2txt.process(file_path)
                        if t and t.strip():
                            return t
                    except Exception:
                        pass
                # Try mammoth
                if MAMMOTH_AVAILABLE:
                    try:
                        with open(file_path, 'rb') as f:
                            result = mammoth.extract_raw_text(f)
                            t = (result.value or '').strip() if result else ''
                            if t:
                                return t
                    except Exception:
                        pass
                # Try OCR on images embedded in docx
                ocr_text = FileProcessor._ocr_docx_images(file_path)
                if ocr_text:
                    return ocr_text

            return extracted if extracted else "[No text extracted from DOCX]"
        except Exception as e:
            return f"[Error reading DOCX: {str(e)}]"

    @staticmethod
    def _read_doc(file_path: str) -> str:
        """Extract text from legacy .doc using Windows Word COM if available, with OCR fallback"""
        if WIN32COM_AVAILABLE:
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(Path(file_path).absolute()))
                try:
                    text = doc.Content.Text
                    if isinstance(text, str) and text.strip():
                        return text
                    # If no text extracted, try alternative method
                    # Try to get text from paragraphs
                    paragraphs = []
                    for para in doc.Paragraphs:
                        para_text = para.Range.Text
                        if para_text and para_text.strip():
                            paragraphs.append(para_text.strip())
                    if paragraphs:
                        return '\n\n'.join(paragraphs)
                    # If no text found via COM, try OCR as fallback
                    ocr_result = FileProcessor._ocr_doc_images(file_path)
                    if ocr_result and ocr_result.strip():
                        return ocr_result
                    return "[No text extracted from .doc]"
                finally:
                    doc.Close(False)
                    word.Quit()
            except Exception as e:
                error_msg = str(e)
                # Try OCR fallback even if COM fails
                try:
                    ocr_result = FileProcessor._ocr_doc_images(file_path)
                    if ocr_result and ocr_result.strip():
                        return ocr_result
                except Exception:
                    pass
                # Return error that will trigger OCR fallback in main.py
                return f"[Error reading DOC via COM: {error_msg}]"
        
        # If win32com not available, try OCR directly
        ocr_result = FileProcessor._ocr_doc_images(file_path)
        if ocr_result and ocr_result.strip():
            return ocr_result
        
        # Fallback guidance - this will trigger OCR attempt in main.py
        return "[Cannot read .doc: install pywin32 (win32com) on Windows or convert to .docx/.pdf]"
    
    @staticmethod
    def _read_excel(file_path: str) -> str:
        """Extract text from Excel file"""
        if PANDAS_AVAILABLE:
            try:
                df = pd.read_excel(file_path, sheet_name=None)
                content_parts = []
                for sheet_name, sheet_df in df.items():
                    content_parts.append(f"Sheet: {sheet_name}\n{sheet_df.to_string()}\n")
                return '\n'.join(content_parts)
            except Exception as e:
                pass
        
        if EXCEL_AVAILABLE:
            try:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                content_parts = []
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    content_parts.append(f"Sheet: {sheet_name}\n")
                    for row in sheet.iter_rows(values_only=True):
                        content_parts.append('\t'.join(str(cell) if cell is not None else '' for cell in row))
                    content_parts.append('\n')
                return '\n'.join(content_parts)
            except Exception as e:
                return f"[Error reading Excel: {str(e)}]"
        
        return "[Excel library not available. Install pandas or openpyxl]"
    
    @staticmethod
    def _read_csv(file_path: str) -> str:
        """Read CSV file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = [','.join(row) for row in reader]
                return '\n'.join(rows)
        except Exception as e:
            return f"[Error reading CSV: {str(e)}]"
    
    @staticmethod
    def _read_json(file_path: str) -> str:
        """Read and format JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            return f"[Error reading JSON: {str(e)}]"
    
    @staticmethod
    def _read_ppt(file_path: str) -> str:
        """Extract text from PowerPoint file"""
        try:
            from ppt_processor import PPTProcessor
            result = PPTProcessor.process_ppt_file(file_path)
            return result.get('slides_text', '[No text extracted from PPT]')
        except ImportError:
            return "[PPT processor not available]"
        except Exception as e:
            return f"[Error reading PPT: {str(e)}]"
    
    @staticmethod
    def process_multiple_files(file_paths: List[str]) -> List[Dict[str, any]]:
        """Process multiple files and return their contents"""
        results = []
        for file_path in file_paths:
            result = FileProcessor.read_file(file_path)
            results.append(result)
        return results



